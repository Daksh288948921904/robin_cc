"""
OSI News — DOCX Report Generator
==================================
Creates a formatted Word (.docx) document from a multi-source
summary, ready for publication or distribution.

Document structure:
  • robin cc masthead
  • Article title (large, bold)
  • Byline + date
  • Body paragraphs
  • Horizontal divider
  • Sources section with numbered list
"""

import io
import re
from datetime import datetime, timezone
from typing import Dict, List

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ────────────────────────────────────────────
_INDIGO   = RGBColor(0x63, 0x66, 0xF1)   # accent #6366f1
_DARK     = RGBColor(0x09, 0x09, 0x0B)   # near-black
_MID_GREY = RGBColor(0x71, 0x71, 0x7A)   # muted text
_BLACK    = RGBColor(0x00, 0x00, 0x00)


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a thin horizontal line."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "6366F1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)


def _set_font(run, bold=False, italic=False, size_pt=11, color: RGBColor = None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _add_para(doc: Document, text: str, bold=False, italic=False,
              size_pt=11, color: RGBColor = None,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6) -> None:
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    run = para.add_run(text)
    _set_font(run, bold=bold, italic=italic, size_pt=size_pt, color=color)


def create_docx(
    summary: Dict,
    coverage: List[Dict],
    main_article: Dict,
) -> io.BytesIO:
    """
    Build a formatted DOCX from a generated summary.

    Args:
        summary:      Output of multi_source_summary.generate_summary()
        coverage:     List of similar coverage article dicts
        main_article: The primary article that was clicked

    Returns:
        BytesIO buffer containing the .docx file contents.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)

    # ── Masthead ──────────────────────────────────────────────
    _add_para(
        doc, "ROBIN CC",
        bold=True, size_pt=9, color=_INDIGO,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2,
    )
    _add_para(
        doc, "Multi-Source News Brief",
        italic=True, size_pt=8, color=_MID_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10,
    )

    _add_horizontal_rule(doc)

    # ── Title ─────────────────────────────────────────────────
    title = summary.get("title") or main_article.get("heading") or "News Summary"
    _add_para(
        doc, title,
        bold=True, size_pt=22, color=_DARK,
        space_before=12, space_after=6,
    )

    # ── Location ──────────────────────────────────────────────
    location = (summary.get("location") or main_article.get("region") or "").strip().upper()
    if location:
        _add_para(doc, f"📍  {location}",
                  bold=False, size_pt=9, color=_MID_GREY,
                  space_before=4, space_after=2)

    # ── Byline row ────────────────────────────────────────────
    byline_raw = summary.get("byline") or f"robin cc | {datetime.now(timezone.utc).strftime('%B %d, %Y')}"
    byline_parts = byline_raw.split("|", 1)
    byline_org  = byline_parts[0].strip()
    byline_date = byline_parts[1].strip() if len(byline_parts) > 1 else datetime.now(timezone.utc).strftime('%B %d, %Y')

    byline_para = doc.add_paragraph()
    byline_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    byline_para.paragraph_format.space_before = Pt(0)
    byline_para.paragraph_format.space_after  = Pt(4)
    run_org = byline_para.add_run(byline_org)
    _set_font(run_org, bold=True, size_pt=9, color=_INDIGO)
    run_sep = byline_para.add_run("  ·  ")
    _set_font(run_sep, size_pt=9, color=_MID_GREY)
    run_date = byline_para.add_run(byline_date)
    _set_font(run_date, italic=True, size_pt=9, color=_MID_GREY)

    # AI BRIEFING badge line
    _add_para(doc, "AI BRIEFING", bold=True, size_pt=8, color=_INDIGO,
              space_before=0, space_after=14)

    _add_horizontal_rule(doc)

    # ── Body — parse sections ─────────────────────────────────
    body = summary.get("body") or ""
    lines = body.split("\n")
    # Group into lede block + named sections
    sections: list = []   # each: {"heading": str|None, "paras": [str]}
    current: dict = {"heading": None, "paras": []}
    for line in lines:
        hm = re.match(r"^##\s+(.+)", line)
        if hm:
            if current["paras"] or current["heading"]:
                sections.append(current)
            current = {"heading": hm.group(1).strip(), "paras": []}
        else:
            t = line.strip()
            if t:
                current["paras"].append(t)
    if current["paras"] or current["heading"]:
        sections.append(current)

    for sec in sections:
        if sec["heading"]:
            # Spacer before section
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(6)
            sp.paragraph_format.space_after  = Pt(0)

            # Section label — uppercase, small, with full-width rule below
            head_para = doc.add_paragraph()
            head_para.paragraph_format.space_before = Pt(0)
            head_para.paragraph_format.space_after  = Pt(0)
            # Full-width bottom border on the heading paragraph
            pPr2 = head_para._p.get_or_add_pPr()
            pBdr2 = OxmlElement("w:pBdr")
            bot2  = OxmlElement("w:bottom")
            bot2.set(qn("w:val"), "single")
            bot2.set(qn("w:sz"), "4")
            bot2.set(qn("w:space"), "2")
            bot2.set(qn("w:color"), "3B3B5C")
            pBdr2.append(bot2)
            pPr2.append(pBdr2)

            run_head = head_para.add_run(sec["heading"].upper())
            run_head.bold = True
            run_head.font.size = Pt(9)
            run_head.font.color.rgb = _INDIGO
            # Approximate letter-spacing via character spacing
            rPr = run_head._r.get_or_add_rPr()
            spacing_el = OxmlElement("w:spacing")
            spacing_el.set(qn("w:val"), "80")   # 8pt spacing ≈ 0.8px
            rPr.append(spacing_el)

            sp2 = doc.add_paragraph()
            sp2.paragraph_format.space_before = Pt(0)
            sp2.paragraph_format.space_after  = Pt(6)

        for para_text in sec["paras"]:
            _add_para(doc, para_text, size_pt=11, color=_BLACK,
                      space_before=0, space_after=8)

    # ── Sources divider ───────────────────────────────────────
    doc.add_paragraph()
    _add_horizontal_rule(doc)

    _add_para(
        doc, "Sources",
        bold=True, size_pt=13, color=_INDIGO,
        space_before=10, space_after=6,
    )

    # Deduplicated sources list
    sources_list = summary.get("sources_list") or []
    if not sources_list:
        # Build from main + coverage
        all_arts = [main_article] + (coverage or [])
        seen: set = set()
        for art in all_arts:
            src = art.get("source_name", "Unknown")
            if src not in seen:
                sources_list.append({
                    "name":     src,
                    "headline": art.get("heading", ""),
                    "url":      art.get("source_url", ""),
                    "region":   art.get("region", ""),
                })
                seen.add(src)

    for n, src in enumerate(sources_list, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(5)
        para.paragraph_format.left_indent  = Cm(0.5)

        # Number + source name
        run_num = para.add_run(f"{n}. ")
        _set_font(run_num, bold=True, size_pt=10, color=_INDIGO)

        run_name = para.add_run(f"{src.get('name','')}  ")
        _set_font(run_name, bold=True, size_pt=10, color=_DARK)

        if src.get("headline"):
            run_hl = para.add_run(f"— {src['headline'][:100]}")
            _set_font(run_hl, size_pt=9, color=_MID_GREY)

        if src.get("url"):
            para2 = doc.add_paragraph()
            para2.paragraph_format.space_before = Pt(0)
            para2.paragraph_format.space_after  = Pt(3)
            para2.paragraph_format.left_indent  = Cm(1)
            run_url = para2.add_run(src["url"])
            _set_font(run_url, size_pt=8, color=_MID_GREY, italic=True)

    # ── Footer ────────────────────────────────────────────────
    doc.add_paragraph()
    _add_horizontal_rule(doc)
    _add_para(
        doc,
        f"Generated by robin cc  ·  {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  ·  "
        f"{len(sources_list)} source(s) synthesised",
        italic=True, size_pt=8, color=_MID_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=0,
    )

    # ── Serialise ─────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
